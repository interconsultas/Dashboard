"""
Genera el documento Word de presentacion del proyecto Interconsultas.
Ejecutar: python docs/generar_documento_entrega.py
Salida:   docs/Interconsultas_Presentacion_Proyecto.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colores de marca (del tailwind.config.ts) ──
NAVY       = RGBColor(0x1A, 0x3F, 0x72)  # #1A3F72
NAVY_DARK  = RGBColor(0x15, 0x33, 0x5C)  # #15335C
GREEN      = RGBColor(0x4E, 0xA2, 0x34)  # #4EA234
GREEN_DARK = RGBColor(0x3D, 0x82, 0x29)  # #3D8229
BLUE_SOFT  = RGBColor(0xE8, 0xEF, 0xF8)  # #E8EFF8
GREEN_SOFT = RGBColor(0xED, 0xF7, 0xE8)  # #EDF7E8
BG_APP     = RGBColor(0xF0, 0xF4, 0xF9)  # #F0F4F9
FOREGROUND = RGBColor(0x1C, 0x2B, 0x3A)  # #1C2B3A
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_600   = RGBColor(0x4B, 0x55, 0x63)
GRAY_400   = RGBColor(0x9C, 0xA3, 0xAF)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH  = os.path.join(SCRIPT_DIR, "..", "dashboard", "public", "logo.png")

doc = Document()

# ── Estilos base ──
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = FOREGROUND
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# ── Margenes ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ──────────────────────────────────────
# Funciones auxiliares
# ──────────────────────────────────────

def set_cell_shading(cell, hex_color):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", sz="4"):
    """Aplica bordes sutiles a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def add_heading_navy(text, level=1):
    """Agrega un encabezado con color navy."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return h


def add_subheading_green(text):
    """Agrega un subtitulo con color verde."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = GREEN
    run.font.bold = True
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(text):
    """Agrega un parrafo de texto normal."""
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_bullet(text, bold_prefix=None):
    """Agrega un item de lista con viñeta."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.color.rgb = NAVY
        run_bold.font.name = "Calibri"
        run_bold.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_highlight_box(text, bg_hex="E8EFF8", border_hex="1A3F72"):
    """Agrega un parrafo destacado con fondo de color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_hex}" w:val="clear"/>')
    pPr.append(shading)

    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="16" w:space="8" w:color="{border_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)

    ind = parse_xml(f'<w:ind {nsdecls("w")} w:left="284" w:right="284"/>')
    pPr.append(ind)

    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = FOREGROUND
    run.font.name = "Calibri"
    return p


def add_branded_table(headers, rows, col_widths=None):
    """Crea una tabla con estilo de la marca."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1A3F72")
        set_cell_borders(cell, "15335C", "4")

    for r_idx, row_data in enumerate(rows):
        bg = "FFFFFF" if r_idx % 2 == 0 else "F0F4F9"
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = FOREGROUND
            set_cell_shading(cell, bg)
            set_cell_borders(cell, "DEE2E6", "4")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_color_swatch(name, hex_code, description):
    """Agrega una muestra de color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)

    run_swatch = p.add_run("  ███  ")
    run_swatch.font.size = Pt(14)
    hex_clean = hex_code.replace("#", "")
    run_swatch.font.color.rgb = RGBColor(
        int(hex_clean[0:2], 16),
        int(hex_clean[2:4], 16),
        int(hex_clean[4:6], 16)
    )

    run_name = p.add_run(f"  {name}")
    run_name.bold = True
    run_name.font.size = Pt(11)
    run_name.font.color.rgb = FOREGROUND
    run_name.font.name = "Calibri"

    run_hex = p.add_run(f"  ({hex_code})")
    run_hex.font.size = Pt(10)
    run_hex.font.color.rgb = GRAY_400
    run_hex.font.name = "Calibri"

    run_desc = p.add_run(f"  —  {description}")
    run_desc.font.size = Pt(10)
    run_desc.font.color.rgb = GRAY_600
    run_desc.font.name = "Calibri"


def add_divider():
    """Agrega una linea divisora verde."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="4EA234"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)


def add_page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════

for _ in range(4):
    doc.add_paragraph()

if os.path.exists(LOGO_PATH):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run().add_picture(LOGO_PATH, width=Cm(6))
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = p_title.add_run("Sistema de Interconsultas")
run_t.font.size = Pt(32)
run_t.font.color.rgb = NAVY
run_t.bold = True
run_t.font.name = "Calibri"

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = p_sub.add_run("Dashboard de Autorizaciones Medicas")
run_s.font.size = Pt(18)
run_s.font.color.rgb = GREEN
run_s.font.name = "Calibri"

doc.add_paragraph()

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = p_desc.add_run("Documentacion de Entrega del Proyecto")
run_d.font.size = Pt(14)
run_d.font.color.rgb = GRAY_600
run_d.font.name = "Calibri"

for _ in range(4):
    doc.add_paragraph()

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_i = p_info.add_run("IPS Manizales  |  2026")
run_i.font.size = Pt(12)
run_i.font.color.rgb = GRAY_400
run_i.font.name = "Calibri"

add_page_break()


# ══════════════════════════════════════════════════
# TABLA DE CONTENIDOS
# ══════════════════════════════════════════════════

add_heading_navy("Contenido", level=1)
add_divider()

toc_items = [
    ("1.", "Que es el Sistema de Interconsultas"),
    ("2.", "Como funciona el sistema"),
    ("3.", "Modulos del Dashboard"),
    ("4.", "Proceso de carga de archivos"),
    ("5.", "Como se garantizan consultas rapidas"),
    ("6.", "Identidad visual y paleta de colores"),
    ("7.", "Infraestructura tecnologica"),
    ("8.", "Seguridad del sistema"),
    ("9.", "Calidad y pruebas automatizadas"),
    ("10.", "Como desplegar el sistema"),
    ("11.", "Resumen de entregables"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r_num = p.add_run(f"{num}  ")
    r_num.font.color.rgb = GREEN
    r_num.bold = True
    r_num.font.size = Pt(12)
    r_num.font.name = "Calibri"
    r_title = p.add_run(title)
    r_title.font.size = Pt(12)
    r_title.font.color.rgb = FOREGROUND
    r_title.font.name = "Calibri"

add_page_break()


# ══════════════════════════════════════════════════
# 1. QUE ES EL SISTEMA
# ══════════════════════════════════════════════════

add_heading_navy("1. Que es el Sistema de Interconsultas", level=1)
add_divider()

add_body(
    "El Sistema de Interconsultas es una plataforma web desarrollada a la medida para "
    "IPS Manizales que permite gestionar y analizar las autorizaciones medicas que la EPS "
    "emite para los afiliados. El sistema reemplaza los procesos manuales de revision de "
    "archivos Excel por un flujo automatizado, seguro y visual."
)

add_highlight_box(
    "En resumen: el sistema recibe los archivos Excel de la EPS, los procesa "
    "automaticamente, detecta errores y duplicados, y presenta los resultados en un "
    "dashboard interactivo con graficas, filtros y exportacion a CSV."
)

add_subheading_green("Que problemas resuelve")

add_bullet(" Elimina el procesamiento manual de archivos Excel con miles de registros", "Automatizacion:  ")
add_bullet(" Detecta automaticamente registros duplicados para evitar datos inflados", "Deduplicacion:  ")
add_bullet(" Cada registro se enriquece con datos del catalogo de profesionales de la IPS", "Cruce de medicos:  ")
add_bullet(" Dashboards interactivos con filtros, graficas de tendencia y rankings", "Visualizacion:  ")
add_bullet(" Descarga de datos filtrados en formato CSV compatible con Excel", "Exportacion:  ")

add_page_break()


# ══════════════════════════════════════════════════
# 2. COMO FUNCIONA
# ══════════════════════════════════════════════════

add_heading_navy("2. Como funciona el sistema", level=1)
add_divider()

add_body(
    "El sistema tiene tres grandes componentes que trabajan juntos de forma automatica:"
)

add_branded_table(
    ["Componente", "Que hace", "Tecnologia"],
    [
        ["Dashboard Web", "Interfaz visual donde los usuarios inician sesion, ven graficas, suben archivos y administran datos", "Next.js 14, React, Tailwind CSS"],
        ["Pipeline ETL", "Motor de procesamiento que lee los archivos Excel, los valida, normaliza y carga en la base de datos", "Python 3.12, pandas"],
        ["Base de Datos", "Almacena todos los registros de autorizaciones, medicos, usuarios y metricas pre-calculadas", "PostgreSQL 16"],
    ],
    col_widths=[3.5, 7, 5.5]
)

add_subheading_green("Flujo de trabajo paso a paso")

steps = [
    ("1. Subir archivo", "El administrador arrastra un archivo .xlsx al area de carga del dashboard."),
    ("2. Validacion automatica", "El sistema verifica la estructura del archivo, detecta columnas faltantes, parsea fechas y numeros, y cruza los datos con el catalogo de medicos."),
    ("3. Informe de previsualizacion", "Se muestra un informe detallado: cuantas filas validas hay, cuantos errores se encontraron, que medicos no se reconocieron y como se distribuyen los estados de autorizacion."),
    ("4. Confirmacion", "El administrador revisa el informe y confirma la carga. Los datos pasan del area temporal a la tabla principal."),
    ("5. Actualizacion del dashboard", "Las vistas pre-calculadas se refrescan automaticamente y el dashboard refleja los nuevos datos de inmediato."),
    ("6. Consulta y analisis", "Todos los usuarios autorizados pueden consultar el dashboard con filtros, ver tendencias mensuales, rankings de profesionales y exportar datos."),
]

for step_title, step_desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    run_num = p.add_run(step_title)
    run_num.bold = True
    run_num.font.color.rgb = NAVY
    run_num.font.name = "Calibri"
    p.add_run(f"\n{step_desc}").font.name = "Calibri"

add_page_break()


# ══════════════════════════════════════════════════
# 3. MODULOS DEL DASHBOARD
# ══════════════════════════════════════════════════

add_heading_navy("3. Modulos del Dashboard", level=1)
add_divider()

add_body(
    "El dashboard esta organizado en modulos accesibles desde un menu lateral (sidebar). "
    "Cada modulo cumple una funcion especifica:"
)

add_subheading_green("3.1 Dashboard Analitico")

add_body(
    "Es la pantalla principal de visualizacion de datos. Ofrece una vista general de "
    "todas las autorizaciones o vistas filtradas por tipo de prestacion."
)

add_branded_table(
    ["Vista", "Que muestra"],
    [
        ["Dashboard General", "Todas las autorizaciones sin filtro de tipo"],
        ["Laboratorios", "Autorizaciones de laboratorio clinico capitado"],
        ["RX (Radiologia)", "Autorizaciones de radiologia diagnostica"],
        ["Ecografias Capitadas", "Ecografias bajo convenio capitado"],
        ["Remisiones Capitadas", "Consultas y procedimientos menores capitados"],
        ["Medicamentos", "Medicamentos PBS y No PBS (excluyendo programas especiales)"],
        ["Remisiones Red Externa", "Consultas derivadas a red externa"],
        ["Procedimientos Dx", "Procedimientos diagnosticos no capitados (actividad, valor agregado)"],
    ],
    col_widths=[4.5, 11.5]
)

add_body("Cada vista incluye:")
add_bullet(" Total de autorizaciones, promedio mensual, profesionales activos, promedio por profesional", "4 tarjetas KPI:  ")
add_bullet(" Comparacion del anio actual vs. el anterior mes a mes", "Grafica de tendencia:  ")
add_bullet(" Profesionales, prestaciones y diagnosticos con mayor volumen", "Top 10 rankings:  ")
add_bullet(" 9 dimensiones de filtrado que se actualizan en cascada", "Filtros interactivos:  ")
add_bullet(" Descarga de hasta 500,000 registros filtrados en formato CSV", "Exportacion:  ")

add_subheading_green("3.2 Carga de Archivos")

add_body(
    "Modulo exclusivo para administradores. Permite subir archivos Excel (.xlsx) de la EPS "
    "y seguir el proceso de validacion en tiempo real."
)
add_bullet(" Arrastrar y soltar archivos (drag & drop)", "Zona de carga:  ")
add_bullet(" Informe detallado de filas, errores, medicos no encontrados", "Previsualizacion:  ")
add_bullet(" Historial de todas las cargas anteriores con estado y metricas", "Historial:  ")

add_subheading_green("3.3 Administracion de Usuarios")

add_body(
    "Permite al administrador crear, editar y desactivar cuentas de usuario del sistema."
)

add_subheading_green("3.4 Administracion de Profesionales")

add_body(
    "Gestion del catalogo de medicos y profesionales de la IPS. Este catalogo es fundamental "
    "porque el ETL lo usa para enriquecer cada autorizacion con datos del medico que la ordeno "
    "(estado activo/inactivo, programa, especialidad, area)."
)

add_page_break()


# ══════════════════════════════════════════════════
# 4. PROCESO DE CARGA
# ══════════════════════════════════════════════════

add_heading_navy("4. Proceso de carga de archivos", level=1)
add_divider()

add_body(
    "El proceso de carga es el corazon operativo del sistema. Transforma archivos Excel "
    "crudos de la EPS en datos estructurados, validados y listos para analisis."
)

add_subheading_green("4.1 Que hace el sistema con cada archivo")

add_branded_table(
    ["Etapa", "Que sucede", "Que se verifica"],
    [
        ["Recepcion", "El archivo se sube al servidor y se le asigna un identificador unico (UUID)", "Extension .xlsx, tamano menor a 150 MB"],
        ["Deteccion de periodo", "Se extrae el mes y anio del nombre del archivo (ej: 202603 = marzo 2026)", "Formato del nombre del archivo"],
        ["Verificacion de duplicados", "Se calcula una huella digital (SHA-256) del archivo completo", "Que no se haya cargado antes el mismo archivo"],
        ["Validacion de estructura", "Se verifican las 56 columnas obligatorias del formato EPS", "Columnas faltantes o sobrantes"],
        ["Limpieza de datos", "Se normalizan textos, fechas y valores numericos", "Fechas invalidas, valores no numericos"],
        ["Cruce de medicos", "Cada registro se enriquece con datos del catalogo de profesionales", "Medicos no encontrados en el catalogo"],
        ["Deduplicacion por fila", "Se calcula una huella unica por registro para evitar duplicados", "Registros ya existentes en la base de datos"],
        ["Carga temporal", "Los datos validados se cargan en un area temporal (staging)", "Integridad de la carga"],
        ["Confirmacion", "El administrador revisa el informe y confirma o cancela", "Decision humana"],
        ["Carga definitiva", "Los datos pasan a la tabla principal y las vistas se actualizan", "Conteo final de insertados vs. duplicados"],
    ],
    col_widths=[3, 6, 5.5]
)

add_subheading_green("4.2 Informe de validacion")

add_body("Despues de procesar el archivo, el sistema muestra un informe detallado que incluye:")

add_bullet(" Cuantas filas tiene el archivo original", "Total de filas:  ")
add_bullet(" Cuantas pasaron todas las validaciones", "Filas validas:  ")
add_bullet(" Cuantas tienen problemas de formato o datos", "Filas con error:  ")
add_bullet(" Cuantas fechas no se pudieron interpretar", "Fechas invalidas:  ")
add_bullet(" Que valores no se pudieron convertir a numero", "Valores invalidos:  ")
add_bullet(" Prestadores que no se encontraron en el catalogo de la IPS", "Medicos no encontrados:  ")
add_bullet(" Que porcentaje de autorizaciones esta en cada estado", "Distribucion de estados:  ")

add_highlight_box(
    "Proteccion contra duplicados: El sistema tiene dos niveles de proteccion. "
    "Primero verifica que el archivo completo no se haya cargado antes (por huella digital). "
    "Luego, al cargar los registros, verifica cada fila individualmente para que no se "
    "dupliquen autorizaciones aunque vengan en archivos diferentes."
)

add_page_break()


# ══════════════════════════════════════════════════
# 5. CONSULTAS RAPIDAS
# ══════════════════════════════════════════════════

add_heading_navy("5. Como se garantizan consultas rapidas", level=1)
add_divider()

add_body(
    "Uno de los principales retos tecnicos del sistema es mantener tiempos de respuesta "
    "rapidos cuando la base de datos contiene millones de registros de autorizaciones. "
    "Para lograr esto, se implementaron cuatro estrategias complementarias:"
)

add_subheading_green("5.1 Vistas materializadas (pre-calculo de datos)")

add_body(
    "En lugar de calcular totales, promedios y conteos cada vez que un usuario abre el "
    "dashboard, el sistema mantiene 8 tablas pre-calculadas que se actualizan "
    "automaticamente despues de cada carga de datos."
)

add_highlight_box(
    "Analogia: Imagine una biblioteca que tiene millones de libros. En lugar de buscar "
    "libro por libro cada vez que alguien pregunta cuantos hay de cada tema, la biblioteca "
    "mantiene un resumen actualizado por tema. Ese resumen es la vista materializada."
)

add_branded_table(
    ["Vista pre-calculada", "Que contiene", "Para que sirve"],
    [
        ["Vista General", "Todas las autorizaciones agrupadas por 12 dimensiones", "Dashboard principal y filtros globales"],
        ["Laboratorios", "Solo autorizaciones de laboratorio clinico", "Dashboard de laboratorios"],
        ["RX", "Solo autorizaciones de radiologia", "Dashboard de radiologia"],
        ["Ecografias", "Solo ecografias capitadas", "Dashboard de ecografias"],
        ["Remisiones Capitadas", "Consultas y procedimientos capitados", "Dashboard de remisiones CAP"],
        ["Medicamentos", "Medicamentos PBS y No PBS", "Dashboard de medicamentos"],
        ["Remisiones Red Externa", "Consultas derivadas a red externa", "Dashboard red externa"],
        ["Procedimientos Dx", "Procedimientos diagnosticos no capitados", "Dashboard procedimientos"],
    ],
    col_widths=[3.5, 6.5, 5.5]
)

add_subheading_green("5.2 Particionamiento de la tabla principal")

add_body(
    "La tabla de autorizaciones esta dividida fisicamente en subtablas mensuales. "
    "Cuando el usuario consulta datos de marzo 2026, la base de datos solo revisa la "
    "subtabla de ese mes, ignorando completamente los datos de otros periodos."
)

add_highlight_box(
    "Analogia: Es como tener un archivador con un cajon por cada mes. "
    "Si necesita los documentos de marzo, solo abre el cajon de marzo en lugar de "
    "revisar todos los cajones del archivador."
)

add_subheading_green("5.3 Indices optimizados")

add_body(
    "Se crearon 14 indices especializados que aceleran las consultas mas frecuentes. "
    "Cada indice esta disenado para un patron de consulta especifico del dashboard:"
)

add_branded_table(
    ["Patron de consulta", "Indice", "Beneficio"],
    [
        ["Filtrar por medico y periodo", "medico + periodo", "Busqueda instantanea de autorizaciones de un medico"],
        ["Distribucion por tipo de prestacion", "periodo + tipo_prestacion", "Graficas de distribucion sin escaneo completo"],
        ["Analisis por diagnostico", "codigo_diagnostico + periodo", "Rankings de diagnosticos con respuesta rapida"],
        ["Vista regional", "regional + periodo + tipo", "Datos filtrados por regional sin retraso"],
        ["Estado de autorizacion", "estado + periodo", "Filtro por estado con acceso directo"],
        ["Programa o especialidad", "programa + periodo", "Analisis por especialidad medica"],
    ],
    col_widths=[4, 4, 7.5]
)

add_subheading_green("5.4 Cache en multiples niveles")

add_body("El sistema almacena temporalmente los resultados de consultas recientes en tres niveles:")

add_branded_table(
    ["Nivel", "Donde", "Duracion", "Que guarda"],
    [
        ["1. Cache del navegador", "En la memoria del navegador del usuario", "Mientras la pestana este abierta", "Respuestas de la API con revalidacion automatica"],
        ["2. Cache HTTP", "Entre el navegador y el servidor", "60 segundos (con revalidacion)", "Encabezados de cache del protocolo HTTP"],
        ["3. Cache del servidor", "En la memoria del servidor Next.js", "5 minutos", "Resultados de consultas a la base de datos"],
    ],
    col_widths=[3, 3.5, 3.5, 6]
)

add_subheading_green("5.5 Filtros en cascada")

add_body(
    "Los filtros del dashboard se actualizan inteligentemente: al seleccionar un filtro, "
    "los demas se ajustan para mostrar solo opciones que tienen datos reales. "
    "Esto evita que el usuario seleccione combinaciones sin resultados y permite que las "
    "consultas sean mas rapidas al reducir el volumen de datos procesados."
)

add_page_break()


# ══════════════════════════════════════════════════
# 6. IDENTIDAD VISUAL
# ══════════════════════════════════════════════════

add_heading_navy("6. Identidad visual y paleta de colores", level=1)
add_divider()

add_body(
    "El dashboard utiliza una paleta de colores profesional basada en los colores "
    "institucionales de la IPS. Todos los componentes visuales — botones, tablas, "
    "graficas, badges, barras laterales — usan consistentemente estos colores."
)

add_subheading_green("6.1 Colores principales")

add_color_swatch("Azul Marino (Navy)", "#1A3F72",
    "Color principal. Sidebar, encabezados, botones primarios, textos destacados")
add_color_swatch("Azul Marino Oscuro", "#15335C",
    "Hover de botones primarios y fondos de enfasis")
add_color_swatch("Verde Institucional", "#4EA234",
    "Color de acento. Botones de exito, badges activos, indicadores positivos")
add_color_swatch("Verde Oscuro", "#3D8229",
    "Hover de botones de confirmacion y exito")

add_subheading_green("6.2 Colores de superficie")

add_color_swatch("Azul Suave", "#E8EFF8",
    "Fondo de badges, hover de botones secundarios, cajas de informacion")
add_color_swatch("Azul Suave 2", "#EEF3FB",
    "Hover de filas en tablas")
add_color_swatch("Verde Suave", "#EDF7E8",
    "Fondo de notificaciones de exito y badges verdes")
add_color_swatch("Fondo de la App", "#F0F4F9",
    "Color de fondo general del shell de la aplicacion")
add_color_swatch("Fondo Alternado", "#F8FAFD",
    "Filas alternadas en tablas (efecto zebra)")

add_subheading_green("6.3 Colores de texto")

add_color_swatch("Texto Principal", "#1C2B3A",
    "Textos de cuerpo, parrafos, datos en tablas")
add_color_swatch("Texto Secundario", "#4B5563",
    "Textos de soporte, descripciones, labels")
add_color_swatch("Texto Terciario", "#9CA3AF",
    "Textos de menor importancia, placeholders, marcas de tiempo")

add_subheading_green("6.4 Tipografia")

add_body(
    "Se utiliza la fuente Inter (de Google Fonts) en los pesos 400, 500, 600, 700 y 800. "
    "Es una tipografia sans-serif disenada especificamente para interfaces digitales, con "
    "excelente legibilidad en pantallas de todas las resoluciones."
)

add_subheading_green("6.5 Aplicacion de colores en la interfaz")

add_branded_table(
    ["Elemento de la interfaz", "Color aplicado", "Motivo"],
    [
        ["Sidebar (barra lateral)", "Azul Marino #1A3F72", "Identidad institucional, transmite profesionalismo"],
        ["Botones principales", "Azul Marino #1A3F72", "Consistencia con la marca, alta visibilidad"],
        ["Botones de confirmacion", "Verde #4EA234", "Color universal de aprobacion y exito"],
        ["Fondo general", "Gris azulado #F0F4F9", "Reduce fatiga visual en uso prolongado"],
        ["Tarjetas KPI", "Blanco con borde suave", "Separacion visual limpia sobre el fondo"],
        ["Badge 'Activo'", "Verde suave #EDF7E8 + texto verde", "Indica estado positivo"],
        ["Badge 'Inactivo'", "Gris suave + texto gris", "Indica estado neutro sin alarma"],
        ["Semaforo de cumplimiento", "Verde / Amarillo / Rojo", "Indicador rapido de metas"],
        ["Tablas", "Encabezado Navy + filas zebra", "Lectura facil de datos tabulares"],
        ["Links del menu activo", "Verde #4EA234", "Indica la seccion actual sin ambiguedad"],
    ],
    col_widths=[4, 4, 7.5]
)

add_page_break()


# ══════════════════════════════════════════════════
# 7. INFRAESTRUCTURA
# ══════════════════════════════════════════════════

add_heading_navy("7. Infraestructura tecnologica", level=1)
add_divider()

add_body(
    "El sistema se despliega como un conjunto de contenedores Docker en un servidor "
    "de DigitalOcean (Droplet). Esta arquitectura garantiza que el sistema funcione de "
    "forma identica en cualquier servidor y se pueda reinstalar en minutos."
)

add_subheading_green("7.1 Componentes de infraestructura")

add_branded_table(
    ["Componente", "Tecnologia", "Funcion"],
    [
        ["Servidor web + ETL", "Node.js 20 + Python 3.12", "Sirve el dashboard y ejecuta el procesamiento de archivos"],
        ["Base de datos", "PostgreSQL 16", "Almacena todos los datos del sistema"],
        ["Contenedores", "Docker + Docker Compose", "Empaqueta y orquesta los servicios"],
        ["Servidor fisico", "DigitalOcean Droplet 4GB", "Infraestructura en la nube ($24/mes)"],
    ],
    col_widths=[3.5, 4, 8]
)

add_subheading_green("7.2 Arquitectura de contenedores")

add_body(
    "El sistema usa dos contenedores que se comunican internamente:"
)

add_bullet(
    " Contiene el dashboard web (Next.js) y el motor de procesamiento ETL (Python). "
    "Expone el puerto 3000 para acceso web.",
    "Contenedor Dashboard:  "
)
add_bullet(
    " Contiene la base de datos PostgreSQL. Los datos se almacenan en un volumen "
    "persistente que sobrevive a reinicios del contenedor.",
    "Contenedor Base de Datos:  "
)

add_highlight_box(
    "El contenedor del dashboard incluye tanto Node.js como Python porque el procesamiento "
    "de archivos Excel se ejecuta como un proceso hijo dentro del mismo contenedor. "
    "Esto simplifica la arquitectura y evita la complejidad de comunicacion entre contenedores separados."
)

add_subheading_green("7.3 Persistencia de datos")

add_body("Los datos del sistema se almacenan en tres volumenes independientes:")

add_branded_table(
    ["Volumen", "Contenido", "Que pasa si se reinicia el servidor"],
    [
        ["pgdata", "Toda la base de datos (tablas, indices, vistas)", "Los datos se conservan intactos"],
        ["uploads", "Archivos Excel subidos temporalmente", "Se conservan los archivos pendientes"],
        ["logs", "Registros detallados de cada procesamiento ETL", "Se conservan los logs historicos"],
    ],
    col_widths=[3, 6, 6.5]
)

add_subheading_green("7.4 Monitoreo automatico (Health Checks)")

add_body(
    "Docker verifica automaticamente que los servicios esten funcionando correctamente. "
    "Si un servicio deja de responder, Docker lo reinicia automaticamente:"
)

add_branded_table(
    ["Servicio", "Verificacion", "Cada cuanto", "Accion si falla"],
    [
        ["Dashboard", "Consulta al endpoint /api/health", "30 segundos", "Reinicio automatico despues de 3 fallos"],
        ["Base de datos", "Comando pg_isready de PostgreSQL", "10 segundos", "Reinicio automatico despues de 5 fallos"],
    ],
    col_widths=[3, 5, 3, 4.5]
)

add_page_break()


# ══════════════════════════════════════════════════
# 8. SEGURIDAD
# ══════════════════════════════════════════════════

add_heading_navy("8. Seguridad del sistema", level=1)
add_divider()

add_body(
    "La seguridad del sistema se implementa en multiples capas para proteger los datos "
    "de los pacientes y garantizar que solo usuarios autorizados accedan a la informacion."
)

add_branded_table(
    ["Capa de seguridad", "Como funciona", "Que protege"],
    [
        ["Autenticacion", "Login con email y contrasena. Las contrasenas se almacenan con hash bcrypt (factor de costo 12)", "Acceso no autorizado al sistema"],
        ["Sesiones JWT", "Token firmado digitalmente que expira automaticamente. Se almacena como cookie HTTP-only", "Suplantacion de identidad"],
        ["Proteccion contra fuerza bruta", "Maximo 5 intentos fallidos por email en 15 minutos", "Ataques de fuerza bruta contra contrasenas"],
        ["Consultas parametrizadas", "Todas las consultas a la BD usan parametros separados, nunca concatenacion de texto", "Inyeccion SQL"],
        ["Validacion de archivos", "Solo se aceptan .xlsx de hasta 150 MB. Los identificadores se validan como UUID", "Archivos maliciosos, ataques de path traversal"],
        ["Gestion de secretos", "Contrasenas de BD y claves de sesion se almacenan en variables de entorno, nunca en el codigo", "Exposicion de credenciales en el repositorio"],
    ],
    col_widths=[3.5, 7, 5]
)

add_page_break()


# ══════════════════════════════════════════════════
# 9. CALIDAD Y PRUEBAS
# ══════════════════════════════════════════════════

add_heading_navy("9. Calidad y pruebas automatizadas", level=1)
add_divider()

add_body(
    "El sistema cuenta con 202 pruebas automatizadas que verifican el correcto "
    "funcionamiento de cada componente. Estas pruebas se pueden ejecutar en cualquier "
    "momento para confirmar que no se han introducido errores."
)

add_branded_table(
    ["Componente", "Cantidad de pruebas", "Que se verifica"],
    [
        ["Validacion de columnas", "10", "Que el sistema detecte correctamente columnas faltantes, sobrantes o vacias"],
        ["Deduplicacion", "12", "Que los hashes se calculen correctamente y detecten archivos/filas duplicadas"],
        ["Generacion de informes", "14", "Que los informes de carga contengan datos correctos y consistentes"],
        ["Deteccion de periodo", "13", "Que el periodo se extraiga correctamente de distintos formatos de nombre"],
        ["Parser de fechas", "14", "Que se interpreten correctamente fechas en formato Excel, ISO y DD/MM/YYYY"],
        ["Cruce de medicos", "9", "Que los medicos se encuentren por usuario_txt y por nombre, y los externos se marquen"],
        ["Preparacion de staging", "7", "Que las columnas se renombren correctamente y los hashes se calculen"],
        ["Conversion de valores", "14", "Que NaN, NaT y otros valores especiales se conviertan correctamente"],
        ["Configuracion de BD", "5", "Que las variables de entorno se lean correctamente"],
        ["Cache del servidor", "9", "Que los datos se almacenen y expiren correctamente"],
        ["Filtros del dashboard", "22", "Que las consultas SQL se construyan correctamente con parametros"],
        ["Registro de vistas", "11", "Que las vistas se resuelvan y validen correctamente"],
        ["Fetcher HTTP", "8", "Que los errores de red se manejen correctamente"],
        ["Spawn del ETL", "6", "Que el proceso Python se lance con los parametros correctos"],
        ["Componentes de UI", "33", "Que los componentes se rendericen correctamente con distintas propiedades"],
        ["Health check", "2", "Que el endpoint responda OK cuando la BD funciona y error cuando no"],
        ["Otros (hooks, helpers)", "13", "Debounce, etiquetas de estado y utilidades menores"],
    ],
    col_widths=[4, 2.5, 9]
)

add_highlight_box(
    "Total: 202 pruebas automatizadas (101 en Python para el ETL + 101 en TypeScript para el dashboard). "
    "Todas las pruebas se ejecutan en menos de 10 segundos."
)

add_page_break()


# ══════════════════════════════════════════════════
# 10. COMO DESPLEGAR
# ══════════════════════════════════════════════════

add_heading_navy("10. Como desplegar el sistema", level=1)
add_divider()

add_body(
    "El despliegue del sistema en un servidor nuevo se realiza en 5 pasos simples. "
    "Todo el proceso toma aproximadamente 10 minutos."
)

add_subheading_green("Requisitos del servidor")

add_branded_table(
    ["Requisito", "Especificacion"],
    [
        ["Proveedor", "DigitalOcean (Droplet)"],
        ["RAM", "4 GB minimo"],
        ["CPU", "2 vCPU"],
        ["Almacenamiento", "80 GB SSD"],
        ["Software", "Docker preinstalado (imagen de Marketplace)"],
        ["Costo mensual", "$24 USD/mes"],
    ],
    col_widths=[4, 12]
)

add_subheading_green("Pasos de despliegue")

deploy_steps = [
    ("Paso 1: Clonar el repositorio",
     "Conectarse al servidor por SSH y descargar el codigo del proyecto desde Git."),
    ("Paso 2: Configurar variables",
     "Copiar el archivo de ejemplo .env.production.example y llenar las contrasenas reales "
     "(base de datos, clave de sesion, URL del servidor)."),
    ("Paso 3: Generar clave de seguridad",
     "Ejecutar un comando para generar una clave aleatoria de 32 caracteres para firmar las sesiones."),
    ("Paso 4: Levantar los servicios",
     "Ejecutar docker-compose up -d para iniciar la base de datos y el dashboard. "
     "La base de datos se inicializa automaticamente con todas las tablas, indices y vistas."),
    ("Paso 5: Verificar",
     "Abrir http://IP_DEL_SERVIDOR:3000 en el navegador. Iniciar sesion con admin@ips.local / Admin2026! "
     "y cambiar la contrasena inmediatamente."),
]

for title, desc in deploy_steps:
    p = doc.add_paragraph()
    run_t = p.add_run(title)
    run_t.bold = True
    run_t.font.color.rgb = NAVY
    run_t.font.name = "Calibri"
    p.add_run(f"\n{desc}").font.name = "Calibri"

add_subheading_green("Operaciones de mantenimiento")

add_branded_table(
    ["Operacion", "Cuando realizarla"],
    [
        ["Ver logs del sistema", "Cuando se necesite diagnosticar un problema"],
        ["Reiniciar un servicio", "Despues de cambiar configuracion o si un servicio no responde"],
        ["Actualizar el codigo", "Cuando se publique una nueva version en el repositorio"],
        ["Hacer backup de la BD", "Semanalmente o antes de cambios importantes"],
        ["Limpiar imagenes Docker viejas", "Mensualmente para liberar espacio en disco"],
    ],
    col_widths=[5, 11]
)

add_page_break()


# ══════════════════════════════════════════════════
# 11. RESUMEN DE ENTREGABLES
# ══════════════════════════════════════════════════

add_heading_navy("11. Resumen de entregables", level=1)
add_divider()

add_body(
    "A continuacion se lista todo lo que se entrega como parte del proyecto:"
)

add_subheading_green("Codigo fuente")

add_branded_table(
    ["Entregable", "Descripcion", "Ubicacion"],
    [
        ["Dashboard Web", "Aplicacion Next.js con 14 endpoints API, 8 vistas de dashboard, 3 modulos de administracion", "dashboard/"],
        ["Pipeline ETL", "Script Python de 1,136 lineas con 10 funciones principales", "etl/"],
        ["Migraciones SQL", "7 archivos SQL con tablas, indices, vistas materializadas y datos iniciales", "etl/sql/"],
        ["Pruebas automatizadas", "202 pruebas (101 Python + 101 TypeScript)", "etl/tests/ + dashboard/src/__tests__/"],
    ],
    col_widths=[3.5, 7, 5]
)

add_subheading_green("Infraestructura")

add_branded_table(
    ["Entregable", "Descripcion", "Ubicacion"],
    [
        ["Dockerfile", "Imagen multi-stage con Node.js 20 + Python 3.12", "Dockerfile"],
        ["Docker Compose", "Orquestacion de 2 servicios (dashboard + PostgreSQL)", "docker-compose.yml"],
        ["Script de inicializacion", "Creacion automatica de BD, usuario, tablas, indices y vistas", "docker/init-db.sh"],
        ["Variables de entorno", "Templates para desarrollo y produccion", ".env.example, .env.production.example"],
        ["Git ignore", "Configuracion para excluir archivos sensibles y temporales", ".gitignore, .dockerignore"],
    ],
    col_widths=[3.5, 7, 5]
)

add_subheading_green("Documentacion")

add_branded_table(
    ["Entregable", "Descripcion", "Ubicacion"],
    [
        ["Documentacion tecnica", "Arquitectura detallada con diagramas, esquemas BD, indices, APIs", "docs/ARQUITECTURA_TECNICA.md"],
        ["Documento de entrega", "Este documento (presentacion amigable del proyecto)", "docs/Interconsultas_Presentacion_Proyecto.docx"],
        ["Contexto para desarrollo", "Guia rapida para desarrolladores que trabajen en el proyecto", "CLAUDE.md"],
    ],
    col_widths=[3.5, 7, 5]
)

doc.add_paragraph()
doc.add_paragraph()

# ── Cierre ──
p_close = doc.add_paragraph()
p_close.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_c = p_close.add_run("Sistema de Interconsultas — IPS Manizales — 2026")
run_c.font.size = Pt(11)
run_c.font.color.rgb = GRAY_400
run_c.font.name = "Calibri"
run_c.italic = True


# ── Guardar ──
output_path = os.path.join(SCRIPT_DIR, "Interconsultas_Presentacion_Proyecto.docx")
doc.save(output_path)
print(f"Documento generado: {output_path}")
